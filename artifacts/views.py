from django.core.paginator import Paginator
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from django.db.models import Q
from .models import Artifact, LoginRecord
from .forms import ArtifactForm
from docx import Document
from django.contrib.staticfiles import finders
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.utils import timezone


# 登录
def user_login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            print(f"用户：{user.username} 登录。")
            login(request, user)
            # 记录登录信息
            ip_address = get_client_ip(request)
            LoginRecord.objects.create(
                user=user,
                login_time=timezone.now(),
                ip_address=ip_address
            )

            return redirect('artifact_list')  # 替换为你的主页路由
        else:
            # 登录失败的处理
            return render(request, 'login.html', {'error': '用户名或密码错误'})
    
    return render(request, 'login.html')

# 注销
def user_logout(request):
    if request.method == 'POST':
        print(f"用户：{request.user.username} 退出。")
        # 更新注销时间
        login_record = LoginRecord.objects.filter(user=request.user, logout_time__isnull=True).first()
        if login_record:
            login_record.logout_time = timezone.now()
            login_record.session_duration = timezone.now() - login_record.login_time
            login_record.save()
        
        logout(request)
        return redirect('login')  # 替换为登录页面路由

    return redirect('login')

# 获取用户 IP 地址
def get_client_ip(request):
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip


# 列出所有文物
@login_required  # 确保用户登录才能访问
def artifact_list(request):
    print(request.user.is_authenticated)  # 打印用户是否已认证
    query = request.GET.get('q')
    if query:
        artifacts = Artifact.objects.filter(
            Q(registration_number__icontains=query) |
            Q(name__icontains=query)
        ).order_by('registration_number')
        LoginRecord.objects.create(
            user=request.user,
            ip_address=get_client_ip(request),
            actions=f'搜索内容：{query}'
        )
    else:
        artifacts = Artifact.objects.all().order_by('registration_number')
    
    # 分页
    paginator = Paginator(artifacts, 30)
    page_number = request.GET.get('page')
    artifacts = paginator.get_page(page_number)
    
    current_user = request.user  # 读取用户信息
    current_user_groups = request.user.groups.values_list('name', flat=True)  # 读取用户组信息

    return render(request, 'artifacts/artifact_list.html', {'artifacts': artifacts, 'query': query, 'current_user': current_user, 'current_user_groups': current_user_groups})

# 创建新文物
@login_required #确保用户登录才能访问
def artifact_create(request):
    if request.method == 'POST':
        form = ArtifactForm(request.POST)
        if form.is_valid():
            form.save()
            LoginRecord.objects.create(
                user=request.user,
                ip_address=get_client_ip(request),
                actions=f'添加文物:{form.name}'
            )
            return redirect('artifact_list')
    else:
        form = ArtifactForm()
    return render(request, 'artifacts/artifact_form.html', {'form': form})

# 更新文物
@login_required #确保用户登录才能访问
def artifact_update(request, pk):
    artifact = get_object_or_404(Artifact, pk=pk)
    if request.method == 'POST':
        form = ArtifactForm(request.POST, instance=artifact)
        if form.is_valid():
            form.save()
            LoginRecord.objects.create(
                user=request.user,
                ip_address=get_client_ip(request),
                actions=f'更新id为：{artifact.id}、名称为：{artifact.name}的文物信息，更新内容为：{request.POST}'
            )
            return redirect('artifact_list')
    else:
        form = ArtifactForm(instance=artifact)
    return render(request, 'artifacts/artifact_form.html', {'form': form})

# 删除文物
@login_required #确保用户登录才能访问
def artifact_delete(request, pk):
    artifact = get_object_or_404(Artifact, pk=pk)
    if request.method == 'POST':
        artifact_name = artifact.name
        artifact.delete()

        LoginRecord.objects.create(
            user=request.user,
            ip_address=get_client_ip(request),
            actions=f'删除id为：{artifact.id}、名称为：{artifact_name}的文物。'
        )
        return redirect('artifact_list')
    return render(request, 'artifacts/artifact_confirm_delete.html', {'artifact': artifact})

# 导出为 WORD 文件
@login_required
def artifact_export_word(request, pk):
    # 获取文物数据
    artifact = Artifact.objects.get(pk=pk)

    # 定义数据字典
    data = {
        "year": artifact.year,
        "month": artifact.month,
        "day": artifact.day,
        "zongdengjihao": artifact.registration_number,
        "fenleihao": artifact.classification_number,
        "name": artifact.name,
        "niandai": artifact.era,
        "jianshu": artifact.quantity,
        "danwei": artifact.unit,
        "chicun": artifact.size,
        "zhongliang": artifact.weight,
        "zhidi": artifact.texture,
        "wancanqingkuang": artifact.condition,
        "laiyuan": artifact.source,
        "ruguanpingzhenghao": artifact.entrance_certificate_number,
        "zhuxiaopingzhenghao": artifact.cancellation_certificate_number,
        "jibie": artifact.level,
        "beizhu": artifact.remarks,
        "fuzeren": artifact.person_in_charge,
        "danganbianhao": artifact.file_number,
        "xingzhuangneirongmiaoshu": artifact.shape_description,
        "dangqianbaocuntiaojian": artifact.current_storage_conditions,
        "mingjitiba": artifact.inscriptions,
    }

    # 指定模板路径
    template_path = finders.find('artifacts/temp.docx')

    # 创建 Word 文档
    doc = Document(template_path)

    # 替换模板中的占位符
    for placeholder, value in data.items():
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # 遍历表格并填充数据（假设表格在文档的第一个位置）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in data.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

    # 创建响应
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="artifact_{pk}.docx"'

    # 保存文档到响应
    doc.save(response)

    LoginRecord.objects.create(
        user=request.user,
        ip_address=get_client_ip(request),
        actions=f'导出id为：{artifact.id}、名称为：{artifact.name}的文物信息为Word文件。'
    )
    return response
